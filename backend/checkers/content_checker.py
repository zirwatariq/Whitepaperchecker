import difflib
import re
from docx import Document
import fitz


def extract_word_paragraphs(docx_path):
    doc = Document(docx_path)
    paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paras.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paras.append(text)
    return paras


def extract_pdf_paragraphs(pdf_path):
    doc = fitz.open(pdf_path)
    paras = []
    for page_num, page in enumerate(doc, start=1):
        for block in page.get_text("blocks"):
            text = block[4].strip()
            if text:
                paras.append({"text": text, "page": page_num})
    doc.close()
    return paras


def normalize(text):
    return re.sub(r'\s+', ' ', text).lower().strip()


def check_content(docx_path, pdf_path):
    word_paras = extract_word_paragraphs(docx_path)
    pdf_paras = extract_pdf_paragraphs(pdf_path)

    word_text = normalize(" ".join(word_paras))
    pdf_text = normalize(" ".join(p["text"] for p in pdf_paras))
    similarity = round(difflib.SequenceMatcher(None, word_text, pdf_text).ratio() * 100, 1)

    # Only compare paragraphs long enough to be meaningful
    word_long = [p for p in word_paras if len(p) > 30]
    pdf_long = pdf_paras  # keep page info

    pdf_norms = [normalize(p["text"]) for p in pdf_long]

    rows = []
    matched_pdf_indices = set()

    for w_text in word_long:
        w_norm = normalize(w_text)
        matches = difflib.get_close_matches(w_norm, pdf_norms, n=1, cutoff=0.82)
        if matches:
            idx = pdf_norms.index(matches[0])
            matched_pdf_indices.add(idx)
            # Only surface as a row if there's a slight but notable difference
            if matches[0] != w_norm:
                rows.append({
                    "status": "mismatch",
                    "word": w_text[:160],
                    "pdf": pdf_long[idx]["text"][:160],
                    "pdf_page": pdf_long[idx]["page"],
                })
        else:
            rows.append({
                "status": "missing",
                "word": w_text[:160],
                "pdf": None,
                "pdf_page": None,
            })

    for i, p in enumerate(pdf_long):
        if i not in matched_pdf_indices and len(p["text"]) > 30:
            rows.append({
                "status": "extra",
                "word": None,
                "pdf": p["text"][:160],
                "pdf_page": p["page"],
            })

    # Sort: missing first, then extra, then mismatches
    order = {"missing": 0, "extra": 1, "mismatch": 2}
    rows.sort(key=lambda r: order.get(r["status"], 3))

    passed = similarity >= 85

    return {
        "passed": passed,
        "similarity_percent": similarity,
        "word_paragraph_count": len(word_paras),
        "pdf_paragraph_count": len(pdf_paras),
        "rows": rows[:25],
    }
